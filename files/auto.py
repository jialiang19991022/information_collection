import pymysql
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


conn = pymysql.connect("34.92.45.21","root","2059598810ASdf_","registration_information",charset="utf8")
cursor = conn.cursor()
cursor.execute("SELECT stu_number FROM collection_information")
result = cursor.fetchall()
for item in result:
    cursor1 = conn.cursor()
    cursor1.execute("SELECT name FROM collection_information where stu_number=%s" % item)
    name = cursor1.fetchone()
    
    cursor2 = conn.cursor()
    cursor2.execute("SELECT phone_number FROM collection_information where stu_number=%s" % item)
    phone_number = cursor2.fetchone()
    
    cursor3 = conn.cursor()
    cursor3.execute("SELECT QQ_number FROM collection_information where stu_number=%s" % item)
    QQ_number = cursor3.fetchone()
    
    cursor4 = conn.cursor()
    cursor4.execute("SELECT first_department FROM collection_information where stu_number=%s" % item)
    first_department = cursor4.fetchone()
    
    cursor5 = conn.cursor()
    cursor5.execute("SELECT second_department FROM collection_information where stu_number=%s" % item)
    second_department = cursor5.fetchone()
    
    cursor6 = conn.cursor()
    cursor6.execute("SELECT notes FROM collection_information where stu_number=%s" % item)
    notes = cursor6.fetchone()
    
    cursor7 = conn.cursor()
    cursor7.execute("SELECT img FROM collection_information where stu_number=%s" % item)
    img = cursor7.fetchone()
    
    cursor8 = conn.cursor()
    cursor8.execute("SELECT sex FROM collection_information where stu_number=%s" % item)
    sex = cursor8.fetchone()
    
    cursor9 = conn.cursor()
    cursor9.execute("SELECT major FROM collection_information where stu_number=%s" % item)
    major = cursor9.fetchone()
    
    cursor10 = conn.cursor()
    cursor10.execute("SELECT experiences FROM collection_information where stu_number=%s" % item)
    experiences = cursor10.fetchone()
    
    cursor11 = conn.cursor()
    cursor11.execute("SELECT wishes FROM collection_information where stu_number=%s" % item)
    wishes = cursor11.fetchone()

    
    doc = Document()
    doc.styles["Normal"].font.name = u'宋体'
    doc.styles["Normal"].font.size = Pt(15)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head = doc.add_heading("", level=1)
    run = head.add_run(u"数科院学生会报名表")
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    tab =doc.add_table(rows=12,cols=2)#创建表格
    tab.style = "Table Grid"
    cell = tab.cell(0,0)
    cell.text = "姓名："
    cell = tab.cell(0,1)
    cell.text = "%s" % name
    
    cell = tab.cell(1,0)
    cell.text = "性别："
    cell = tab.cell(1,1)
    cell.text = "%s" % sex
    
    cell = tab.cell(2,0)
    cell.text = "学号："
    cell = tab.cell(2,1)
    cell.text = "%s" % item
    
    cell = tab.cell(3,0)
    cell.text = "专业："
    cell = tab.cell(3,1)
    cell.text = "%s" % major
    
    cell = tab.cell(4,0)
    cell.text = "手机号码："
    cell = tab.cell(4,1)
    cell.text = "%s" % phone_number
    
    cell = tab.cell(5,0)
    cell.text = "QQ号码："
    cell = tab.cell(5,1)
    cell.text = "%s" % QQ_number
    
    cell = tab.cell(6,0)
    cell.text = "第一意向部门："
    cell = tab.cell(6,1)
    cell.text = "%s" % first_department
    
    cell = tab.cell(7,0)
    cell.text = "第二意向部门："
    cell = tab.cell(7,1)
    cell.text = "%s" % second_department
    
    cell = tab.cell(8,0)
    cell.text = "自我介绍："
    cell = tab.cell(8,1)
    cell.text = "%s" % notes
    
    cell = tab.cell(9,0)
    cell.text = "工作经历："
    cell = tab.cell(9,1)
    cell.text = "%s" % experiences
    
    cell = tab.cell(10,0)
    cell.text = "愿景期待："
    cell = tab.cell(10,1)
    cell.text = "%s" % wishes
    
    cell = tab.cell(11,0)
    cell.text = "照片："
    cell = tab.cell(11,1)
    
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("%s" % img,width=Inches(1.5),height=Inches(1.5))

    
    doc.save("%s.docx" % item)

